#!/usr/bin/env python3
"""
make_cheatsheet.py — convert cheatsheet.md to a formatted .docx

Usage:
    python make_cheatsheet.py [input.md] [output.docx]

Defaults:
    input  — cheatsheet.md  (same directory as this script)
    output — rest_fetcher_cheatsheet.docx (same directory as this script)

Dependencies:
    pip install python-docx mistune
"""

import html
import sys
from pathlib import Path

import mistune
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Twips

# ---------------------------------------------------------------------------
# Page geometry (US Letter, 1" margins)
# ---------------------------------------------------------------------------
PAGE_WIDTH_DXA = 12240
MARGIN_DXA = 1440
CONTENT_DXA = PAGE_WIDTH_DXA - 2 * MARGIN_DXA  # 9360

# ---------------------------------------------------------------------------
# Colour palette  (matches original docx)
# ---------------------------------------------------------------------------
C_HEADING_1 = RGBColor(0x1A, 0x3A, 0x5C)  # dark navy
C_HEADING_2 = RGBColor(0x1A, 0x3A, 0x5C)
C_HEADING_3 = RGBColor(0x2E, 0x75, 0xB6)  # medium blue
C_BODY = RGBColor(0x33, 0x33, 0x33)  # near-black
C_CODE_FG = RGBColor(0x1A, 0x3A, 0x5C)  # navy (inline code)
C_CODE_BG = RGBColor(0xF2, 0xF6, 0xFA)  # light blue-grey (code blocks)
C_NOTE_BG = RGBColor(0xFF, 0xF8, 0xE7)  # pale amber (blockquotes)
C_NOTE_BAR = RGBColor(0xE8, 0xB8, 0x30)  # amber bar
C_TH_BG = RGBColor(0x1A, 0x3A, 0x5C)  # table header bg
C_TH_FG = RGBColor(0xFF, 0xFF, 0xFF)  # table header text
C_TD_BG_ALT = RGBColor(0xF7, 0xFA, 0xFD)  # table row alt bg
C_BORDER = RGBColor(0xBF, 0xCF, 0xE0)  # table border
C_HR = RGBColor(0xBF, 0xCF, 0xE0)  # horizontal rule

# ---------------------------------------------------------------------------
# Font sizes
# ---------------------------------------------------------------------------
SZ_H1 = Pt(16)
SZ_H2 = Pt(13)
SZ_H3 = Pt(11)
SZ_BODY = Pt(9.5)
SZ_CODE = Pt(8.5)
SZ_NOTE = Pt(9)

FONT_BODY = 'Calibri'
FONT_CODE = 'Courier New'

# ---------------------------------------------------------------------------
# AST compatibility helpers
# ---------------------------------------------------------------------------


def _node_type(node) -> str | None:
    return node.get('type') if isinstance(node, dict) else None


def _node_text(node, default: str = '') -> str:
    if node is None:
        return default
    if isinstance(node, str):
        return node
    if not isinstance(node, dict):
        return default
    for key in ('raw', 'text', 'content'):
        val = node.get(key)
        if val is not None:
            return str(val)
    return default


def _children(node):
    if node is None:
        return []
    if isinstance(node, str):
        return [node]
    if isinstance(node, (list, tuple)):
        return list(node)
    if not isinstance(node, dict):
        return []
    children = node.get('children')
    if children is None:
        text = _node_text(node, '')
        return [text] if text else []
    if isinstance(children, (list, tuple)):
        return list(children)
    return [children]


def _first_contentful(items):
    for item in items or []:
        if item is None:
            continue
        if isinstance(item, str) and item.strip():
            return item
        if isinstance(item, dict):
            if _node_type(item) != 'blank_line':
                if _children(item) or _node_text(item, '').strip() or _node_type(item):
                    return item
    return None


# ---------------------------------------------------------------------------
# Table column-width heuristics
# ---------------------------------------------------------------------------


def _classify_table(head_texts: list[str], body_rows: list[list]) -> list[int]:
    """
    Return a list of column widths in DXA that sum to CONTENT_DXA.

    Strategy: if the table looks like a key/description table (2 columns,
    first column header is one of the known key-column names, or first column
    cells are mostly backtick-quoted), give the first column ~28% and the
    rest to description.  For 3-column tables we split roughly evenly but
    give the last (widest) column more room.  Everything else splits evenly.
    """
    n = len(head_texts)
    if n == 0:
        return [CONTENT_DXA]

    KEY_HEADERS = {'key', 'type', 'mode', 'exception', 'callback', 'method', 'component'}

    if n == 2:
        h0 = head_texts[0].strip().lower()
        # check first-column header name
        is_key_col = h0 in KEY_HEADERS
        if not is_key_col and body_rows:
            # check whether ≥60% of first-column cells are code-only
            code_count = sum(1 for row in body_rows if row and _is_code_only(row[0]))
            is_key_col = code_count / len(body_rows) >= 0.6
        if is_key_col:
            w0 = int(CONTENT_DXA * 0.28)
            return [w0, CONTENT_DXA - w0]
        else:
            half = CONTENT_DXA // 2
            return [half, CONTENT_DXA - half]

    if n == 3:
        w0 = int(CONTENT_DXA * 0.18)
        w1 = int(CONTENT_DXA * 0.22)
        w2 = CONTENT_DXA - w0 - w1
        return [w0, w1, w2]

    # fallback: equal columns
    w = CONTENT_DXA // n
    widths = [w] * n
    widths[-1] = CONTENT_DXA - w * (n - 1)
    return widths


def _is_code_only(cell_children) -> bool:
    """True if the cell contains only a single inline-code node.

    Mistune AST schemas vary:
      - some represent cell children as a list of dict nodes
      - others may yield raw strings (or a mix of strings and dicts)
      - occasionally a single dict is used instead of a list
    This helper normalizes those cases so table sizing doesn't crash.
    """
    if cell_children is None:
        return False

    # Normalize to a list
    if isinstance(cell_children, dict):
        items = [cell_children]
    elif isinstance(cell_children, (list, tuple)):
        items = list(cell_children)
    elif isinstance(cell_children, str):
        # raw text cell -> not code-only
        return False
    else:
        return False

    norm = []
    for c in items:
        if isinstance(c, str):
            # treat as plain text node
            if c.strip():
                norm.append({'type': 'text', 'raw': c})
            continue
        if isinstance(c, dict):
            # skip explicit blank_line nodes
            if c.get('type') == 'blank_line':
                continue
            # skip empty text-ish nodes
            raw = c.get('raw') or c.get('text') or ''
            if c.get('type') in ('text', 'linebreak') and not str(raw).strip():
                continue
            norm.append(c)

    if len(norm) != 1:
        return False

    t = norm[0].get('type')
    return t in ('codespan', 'inline_code', 'code_inline')


# ---------------------------------------------------------------------------
# XML helpers
# ---------------------------------------------------------------------------


def _get_or_replace(parent, tag: str):
    """Return existing child with *tag* or create and append a new one."""
    existing = parent.find(qn(tag))
    if existing is not None:
        parent.remove(existing)
    el = OxmlElement(tag)
    return el


# tcPr child order (OOXML CT_TcPr):
#   cnfStyle, tcW, gridSpan, hMerge, vMerge, tcBorders, shd, noWrap, tcMar, ...
_TCP_ORDER = [
    'w:cnfStyle',
    'w:tcW',
    'w:gridSpan',
    'w:hMerge',
    'w:vMerge',
    'w:tcBorders',
    'w:shd',
    'w:noWrap',
    'w:tcMar',
    'w:textDirection',
    'w:tcFitText',
    'w:vAlign',
]

# pPr child order (OOXML CT_PPr, simplified):
#   pStyle, keepNext, keepLines, pageBreakBefore, framePr,
#   pBdr, shd, tabs, spacing, ind, contextualSpacing, ...
_PPR_ORDER = [
    'w:pStyle',
    'w:keepNext',
    'w:keepLines',
    'w:pageBreakBefore',
    'w:framePr',
    'w:suppressLineNumbers',
    'w:pBdr',
    'w:shd',
    'w:tabs',
    'w:suppressAutoHyphens',
    'w:spacing',
    'w:ind',
    'w:contextualSpacing',
    'w:adjustRightInd',
    'w:jc',
]


def _insert_ordered(parent, el, order: list[str]):
    """Insert *el* into *parent* respecting *order*.  Replaces if tag exists."""
    tag = el.tag if '}' not in el.tag else 'w:' + el.tag.split('}')[1]
    # remove existing element with same tag
    existing = parent.find(qn(tag))
    if existing is not None:
        parent.remove(existing)
    # find insertion point
    try:
        el_idx = order.index(tag)
    except ValueError:
        parent.append(el)
        return
    # scan existing children for the first one that comes after el in order
    insert_before = None
    for child in parent:
        ctag = child.tag if '}' not in child.tag else 'w:' + child.tag.split('}')[1]
        try:
            if order.index(ctag) > el_idx:
                insert_before = child
                break
        except ValueError:
            pass
    if insert_before is not None:
        insert_before.addprevious(el)
    else:
        parent.append(el)


def _set_cell_width(cell, width_dxa: int):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:type'), 'dxa')
    tcW.set(qn('w:w'), str(width_dxa))
    _insert_ordered(tcPr, tcW, _TCP_ORDER)


def _set_cell_borders(cell, color_hex: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top', 'start', 'bottom', 'end'):  # strict OOXML: start/end not left/right
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '1')
        el.set(qn('w:color'), color_hex)
        tcBorders.append(el)
    _insert_ordered(tcPr, tcBorders, _TCP_ORDER)


def _set_cell_shading(cell, fill_hex: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    _insert_ordered(tcPr, shd, _TCP_ORDER)


def _set_cell_margins(cell, top=80, bottom=80, left=120, right=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side, val in (
        ('top', top),
        ('left', left),  # transitional: left/right not start/end
        ('bottom', bottom),
        ('right', right),
    ):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:type'), 'dxa')
        el.set(qn('w:w'), str(val))
        tcMar.append(el)
    _insert_ordered(tcPr, tcMar, _TCP_ORDER)


def _rgb_hex(color: RGBColor) -> str:
    return f'{color[0]:02X}{color[1]:02X}{color[2]:02X}'


def _add_bottom_border_to_para(para, color: RGBColor, size: int = 6):
    pPr = para._p.get_or_add_pPr()
    pBdr = pPr.find(qn('w:pBdr'))
    if pBdr is None:
        pBdr = OxmlElement('w:pBdr')
        _insert_ordered(pPr, pBdr, _PPR_ORDER)
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single')
    bot.set(qn('w:sz'), str(size))
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), _rgb_hex(color))
    pBdr.append(bot)  # bottom comes after top and start in CT_PBdr


def _add_left_bar_to_para(para, color: RGBColor, size: int = 18):
    pPr = para._p.get_or_add_pPr()
    pBdr = pPr.find(qn('w:pBdr'))
    if pBdr is None:
        pBdr = OxmlElement('w:pBdr')
        _insert_ordered(pPr, pBdr, _PPR_ORDER)
    left = OxmlElement('w:left')  # transitional OOXML uses left (not start) for pBdr
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), str(size))
    left.set(qn('w:space'), '4')
    left.set(qn('w:color'), _rgb_hex(color))
    pBdr.append(left)


def _shade_para_bg(para, color: RGBColor):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), _rgb_hex(color))
    _insert_ordered(pPr, shd, _PPR_ORDER)


def _set_para_indent(para, left_twips: int, hanging_twips: int = 0):
    """Set paragraph indents in *twips* (1/20 pt), not EMUs.

    Important: python-docx Length objects such as Twips(360) stringify to EMUs
    (e.g. 228600), but OOXML <w:ind> expects raw twip counts. Passing a
    Length object straight through makes Word interpret the indent as hundreds
    of times larger than intended, which can cascade into pagination/layout
    problems on Word for Windows.
    """
    pPr = para._p.get_or_add_pPr()
    ind = OxmlElement('w:ind')
    ind.set(qn('w:left'), str(int(left_twips)))
    if hanging_twips:
        ind.set(qn('w:hanging'), str(int(hanging_twips)))
    _insert_ordered(pPr, ind, _PPR_ORDER)


def _set_para_spacing(para, before: int = 0, after: int = 0, line: int | None = None):
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), str(before))
    spacing.set(qn('w:after'), str(after))
    if line is not None:
        spacing.set(qn('w:line'), str(line))
        spacing.set(qn('w:lineRule'), 'auto')
    _insert_ordered(pPr, spacing, _PPR_ORDER)


# ---------------------------------------------------------------------------
# Run-level inline rendering
# ---------------------------------------------------------------------------


def _render_inline(
    para,
    nodes: list,
    bold: bool = False,
    italic: bool = False,
    base_size: Pt = SZ_BODY,
    base_color: RGBColor = C_BODY,
):
    """Recursively add runs to *para* from an inline node list.

    Mistune AST differs between versions/plugins:
      - inline nodes may be dicts, strings, or even nested lists
      - text payload may live in keys like 'raw', 'text', or 'content'
    This renderer is intentionally tolerant so we don't silently drop text.
    """
    if nodes is None:
        return

    # Normalize single inline node into a list
    if isinstance(nodes, (dict, str)):
        nodes = [nodes]

    for node in nodes:
        # Normalize node types
        if node is None:
            continue
        if isinstance(node, str):
            if node:
                run = para.add_run(_decode_text(node))
                _style_run(run, bold=bold, italic=italic, size=base_size, color=base_color)
            continue
        if isinstance(node, (list, tuple)):
            _render_inline(
                para,
                list(node),
                bold=bold,
                italic=italic,
                base_size=base_size,
                base_color=base_color,
            )
            continue
        if not isinstance(node, dict):
            continue

        t = _node_type(node)

        # text-like payload helper
        raw_text = _node_text(node, '')

        if t in ('text', 'raw_text'):
            if raw_text:
                run = para.add_run(_decode_text(raw_text))
                _style_run(run, bold=bold, italic=italic, size=base_size, color=base_color)

        elif t == 'softbreak':
            run = para.add_run(' ')
            _style_run(run, size=base_size, color=base_color)

        elif t == 'linebreak':
            br_run = OxmlElement('w:r')
            br = OxmlElement('w:br')
            br_run.append(br)
            para._p.append(br_run)

        elif t in ('codespan', 'inline_code', 'code_inline'):
            if raw_text:
                run = para.add_run(_decode_text(raw_text))
                run.font.name = FONT_CODE
                run.font.size = Pt(base_size.pt - 0.5)
                run.font.color.rgb = C_CODE_FG
                run.font.bold = bold

        elif t == 'strong':
            _render_inline(
                para,
                _children(node),
                bold=True,
                italic=italic,
                base_size=base_size,
                base_color=base_color,
            )

        elif t == 'emphasis':
            _render_inline(
                para,
                _children(node),
                bold=bold,
                italic=True,
                base_size=base_size,
                base_color=base_color,
            )

        elif t == 'link':
            # render link text only (no hyperlink machinery needed for cheatsheet)
            _render_inline(
                para,
                _children(node),
                bold=bold,
                italic=italic,
                base_size=base_size,
                base_color=base_color,
            )

        elif t == 'blank_line':
            pass

        else:
            # unknown inline — render children if present, else raw/text content
            children = _children(node)
            if children:
                _render_inline(
                    para,
                    children,
                    bold=bold,
                    italic=italic,
                    base_size=base_size,
                    base_color=base_color,
                )
            elif raw_text:
                run = para.add_run(_decode_text(raw_text))
                _style_run(run, bold=bold, italic=italic, size=base_size, color=base_color)


def _decode_text(text: str) -> str:
    """Decode HTML entities that may be emitted literally by markdown parsing.

    This keeps content like ``&nbsp;`` from showing up verbatim in Word.
    We intentionally preserve the decoded Unicode characters (e.g. NBSP)
    instead of collapsing them to regular spaces.
    """
    if text is None:
        return ''
    return html.unescape(str(text))


def _style_run(run, bold=False, italic=False, size: Pt = SZ_BODY, color: RGBColor = C_BODY):
    run.font.name = FONT_BODY
    run.font.size = size
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color


# ---------------------------------------------------------------------------
# Block-level rendering
# ---------------------------------------------------------------------------


class DocxRenderer:
    def __init__(self, doc: Document):
        self.doc = doc

    # -- helpers -------------------------------------------------------------

    @staticmethod
    def _attrs(node: dict) -> dict:
        """Return node attributes across Mistune AST variants."""
        attrs = node.get('attrs') if isinstance(node, dict) else None
        return attrs if isinstance(attrs, dict) else {}

    @staticmethod
    def _children_of(node):
        return _children(node)

    @staticmethod
    def _text_of(node, default: str = '') -> str:
        return _node_text(node, default)

    @classmethod
    def _get_level(cls, node: dict, default: int = 1) -> int:
        attrs = cls._attrs(node)
        level = attrs.get('level', node.get('level', default))
        try:
            return int(level)
        except (TypeError, ValueError):
            return default

    @classmethod
    def _is_ordered_list(cls, node: dict) -> bool:
        attrs = cls._attrs(node)
        val = attrs.get('ordered', node.get('ordered', False))
        return bool(val)

    def _para(
        self,
        text: str = '',
        bold: bool = False,
        size: Pt = SZ_BODY,
        color: RGBColor = C_BODY,
        align=WD_ALIGN_PARAGRAPH.LEFT,
    ) -> object:
        p = self.doc.add_paragraph()
        p.alignment = align
        if text:
            run = p.add_run(text)
            _style_run(run, bold=bold, size=size, color=color)
        return p

    def _inline_para(
        self,
        nodes: list,
        size: Pt = SZ_BODY,
        color: RGBColor = C_BODY,
        align=WD_ALIGN_PARAGRAPH.LEFT,
    ) -> object:
        p = self.doc.add_paragraph()
        p.alignment = align
        _render_inline(p, nodes, base_size=size, base_color=color)
        return p

    # -- block handlers ------------------------------------------------------

    def render_heading(self, node: dict):
        level = self._get_level(node)
        children = self._children_of(node)

        if level == 1:
            p = self._inline_para(children, size=SZ_H1, color=C_HEADING_1)
            p.runs[0].font.bold = True if p.runs else None
            for run in p.runs:
                run.font.bold = True
            _set_para_spacing(p, before=0, after=120)
            _add_bottom_border_to_para(p, C_HEADING_1, size=8)

        elif level == 2:
            p = self._inline_para(children, size=SZ_H2, color=C_HEADING_2)
            for run in p.runs:
                run.font.bold = True
            _set_para_spacing(p, before=240, after=80)
            _add_bottom_border_to_para(p, C_HR, size=4)

        elif level == 3:
            p = self._inline_para(children, size=SZ_H3, color=C_HEADING_3)
            for run in p.runs:
                run.font.bold = True
            _set_para_spacing(p, before=160, after=60)

        else:
            p = self._inline_para(children, size=SZ_BODY, color=C_HEADING_3)
            for run in p.runs:
                run.font.bold = True
                run.font.italic = True
            _set_para_spacing(p, before=120, after=40)

    def render_paragraph(self, node: dict):
        children = self._children_of(node)
        p = self._inline_para(children, size=SZ_BODY, color=C_BODY)
        _set_para_spacing(p, before=0, after=80)

    def render_block_code(self, node: dict):
        code = self._text_of(node, '').rstrip('\n')
        lines = code.split('\n')
        p = self.doc.add_paragraph()
        _set_para_spacing(p, before=60, after=60)
        _shade_para_bg(p, C_CODE_BG)
        _set_para_indent(p, left_twips=180)
        for i, line in enumerate(lines):
            if i > 0:
                # w:br must be inside a w:r element
                br_run = OxmlElement('w:r')
                br = OxmlElement('w:br')
                br_run.append(br)
                p._p.append(br_run)
            run = p.add_run(line)
            run.font.name = FONT_CODE
            run.font.size = SZ_CODE
            run.font.color.rgb = C_CODE_FG

    def render_block_quote(self, node: dict):
        children = self._children_of(node)
        for child in children:
            if child.get('type') == 'paragraph':
                p = self._inline_para(self._children_of(child), size=SZ_NOTE, color=C_BODY)
                _shade_para_bg(p, C_NOTE_BG)
                _add_left_bar_to_para(p, C_NOTE_BAR)
                _set_para_indent(p, left_twips=360)
                _set_para_spacing(p, before=40, after=40)
            else:
                self.render_node(child)

    def render_thematic_break(self, _node: dict):
        p = self.doc.add_paragraph()
        _set_para_spacing(p, before=80, after=80)
        _add_bottom_border_to_para(p, C_HR, size=4)

    def render_list(self, node: dict, level: int = 0):
        ordered = self._is_ordered_list(node)
        items = self._children_of(node)
        for idx, item in enumerate(items, start=1):
            self._render_list_item(item, ordered=ordered, number=idx, level=level)

    def _render_list_item(self, item: dict, ordered: bool, number: int, level: int):
        children = self._children_of(item)
        first = _first_contentful(children)

        inline_nodes = []
        remainder = []
        if isinstance(first, dict) and _node_type(first) in (
            'paragraph',
            'block_text',
            'list_item',
        ):
            inline_nodes = self._children_of(first)
            seen_first = False
            for child in children:
                if not seen_first and child is first:
                    seen_first = True
                    continue
                remainder.append(child)
        elif first is not None and _node_type(first) != 'list':
            inline_nodes = [first]
            seen_first = False
            for child in children:
                if not seen_first and child is first:
                    seen_first = True
                    continue
                remainder.append(child)
        else:
            remainder = list(children)

        p = self.doc.add_paragraph()
        _set_para_indent(p, left_twips=360 + level * 360, hanging_twips=240)
        _set_para_spacing(p, before=0, after=40)

        bullet = f'{number}.' if ordered else '•'
        run = p.add_run(bullet + '\t')
        _style_run(run, size=SZ_BODY, color=C_BODY)

        _render_inline(p, inline_nodes, base_size=SZ_BODY, base_color=C_BODY)

        for child in remainder:
            ctype = _node_type(child)
            if ctype == 'list':
                self.render_list(child, level=level + 1)
            elif ctype == 'paragraph':
                p2 = self._inline_para(self._children_of(child), size=SZ_BODY, color=C_BODY)
                _set_para_indent(p2, left_twips=360 + level * 360 + 240)
                _set_para_spacing(p2, before=0, after=40)
            elif isinstance(child, dict):
                self.render_node(child)

    def render_table(self, node: dict):
        children = self._children_of(node)
        head_node = next((c for c in children if c['type'] == 'table_head'), None)
        body_node = next((c for c in children if c['type'] == 'table_body'), None)

        head_cells = self._children_of(head_node) if head_node else []
        body_rows = []
        if body_node:
            for row in self._children_of(body_node):
                body_rows.append(self._children_of(row))

        n_cols = len(head_cells)
        if n_cols == 0:
            return

        # collect plain text from header for heuristic
        head_texts = [_inline_to_text(self._children_of(c)) for c in head_cells]
        col_widths = _classify_table(head_texts, body_rows)

        n_rows = 1 + len(body_rows)
        table = self.doc.add_table(rows=n_rows, cols=n_cols)
        table.style = 'Table Grid'

        # set table width and lock column widths (fixed layout)
        tbl = table._tbl
        tblPr = tbl.tblPr
        tblW = tblPr.find(qn('w:tblW'))
        if tblW is None:
            tblW = OxmlElement('w:tblW')
            tblPr.append(tblW)
        tblW.set(qn('w:type'), 'dxa')
        tblW.set(qn('w:w'), str(CONTENT_DXA))

        # tblLayout fixed — critical: tells Word to honour explicit col widths
        # insert before tblLook (which python-docx places last)
        tblLayout = OxmlElement('w:tblLayout')
        tblLayout.set(qn('w:type'), 'fixed')
        tblLook = tblPr.find(qn('w:tblLook'))
        if tblLook is not None:
            tblLook.addprevious(tblLayout)
        else:
            tblPr.append(tblLayout)

        # disable Word's autofit behaviour
        tbl.allow_autofit = False

        # header row
        hrow = table.rows[0]
        for ci, (cell_node, width) in enumerate(zip(head_cells, col_widths, strict=False)):
            cell = hrow.cells[ci]
            _set_cell_width(cell, width)
            _set_cell_shading(cell, _rgb_hex(C_TH_BG))
            _set_cell_borders(cell, _rgb_hex(C_BORDER))
            _set_cell_margins(cell)

            # clear default paragraph, add styled one
            cell.paragraphs[0]._element.getparent().remove(cell.paragraphs[0]._element)
            p = cell.add_paragraph()
            _set_para_spacing(p, before=0, after=0)
            _render_inline(
                p, self._children_of(cell_node), bold=True, base_size=SZ_BODY, base_color=C_TH_FG
            )

        # body rows
        for ri, row_cells in enumerate(body_rows):
            drow = table.rows[ri + 1]
            bg_hex = _rgb_hex(C_TD_BG_ALT) if ri % 2 else 'FFFFFF'
            for ci in range(n_cols):
                cell = drow.cells[ci]
                cell_node = row_cells[ci] if ci < len(row_cells) else {}
                cell_nodes = self._children_of(cell_node)
                width = col_widths[ci]

                _set_cell_width(cell, width)
                _set_cell_shading(cell, bg_hex)
                _set_cell_borders(cell, _rgb_hex(C_BORDER))
                _set_cell_margins(cell)

                cell.paragraphs[0]._element.getparent().remove(cell.paragraphs[0]._element)
                p = cell.add_paragraph()
                _set_para_spacing(p, before=0, after=0)
                _render_inline(p, cell_nodes, base_size=SZ_BODY, base_color=C_BODY)

        # spacing after table via tblPr tblCellMar or just a small para
        p = self.doc.add_paragraph()
        _set_para_spacing(p, before=0, after=60)

    # -- dispatcher ----------------------------------------------------------

    def render_node(self, node: dict):
        # tolerate nodes arriving as strings/lists from different parser versions
        if node is None:
            return
        if isinstance(node, str):
            # treat stray strings as a paragraph
            p = self._para(node, size=SZ_BODY, color=C_BODY)
            _set_para_spacing(p, before=0, after=80)
            return
        if isinstance(node, (list, tuple)):
            for ch in node:
                self.render_node(ch)
            return
        if not isinstance(node, dict):
            return

        t = node.get('type')

        if t == 'blank_line':
            return

        # Headings / paragraphs
        if t == 'heading':
            self.render_heading(node)
            return
        if t == 'paragraph':
            self.render_paragraph(node)
            return

        # Block code: mistune versions vary ('block_code', 'fenced_code', 'code_block', etc.)
        if t == 'block_code' or (
            isinstance(t, str)
            and ('code' in t)
            and t not in ('codespan', 'inline_code', 'code_inline')
        ):
            # Only treat it as a block if it has any direct text payload; otherwise fall through.
            if (
                self._text_of(node, '').strip()
                or node.get('info') is not None
                or node.get('lang') is not None
            ):
                self.render_block_code(node)
                return

        if t == 'block_quote':
            self.render_block_quote(node)
            return
        if t == 'thematic_break':
            self.render_thematic_break(node)
            return
        if t == 'list':
            self.render_list(node)
            return
        if t == 'table':
            self.render_table(node)
            return

        # Fallback: render children if present, otherwise render any text payload
        children = self._children_of(node)
        if children:
            for child in children:
                self.render_node(child)
        else:
            txt = self._text_of(node, '').strip()
            if txt:
                p = self._para(txt, size=SZ_BODY, color=C_BODY)
                _set_para_spacing(p, before=0, after=80)

    def render_all(self, ast: list):
        for node in ast:
            self.render_node(node)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _inline_to_text(nodes) -> str:
    """Extract plain text from inline-ish structures for heuristics."""
    parts = []
    for n in _children(nodes):
        if isinstance(n, str):
            parts.append(n)
            continue
        if not isinstance(n, dict):
            continue
        text = _node_text(n, '')
        if text:
            parts.append(text)
        kids = _children(n)
        if kids and kids != [text]:
            parts.append(_inline_to_text(kids))
    return ''.join(parts)


# ---------------------------------------------------------------------------
# Document setup
# ---------------------------------------------------------------------------


def _setup_document() -> Document:
    doc = Document()

    # page size and margins
    section = doc.sections[0]
    section.page_width = Twips(PAGE_WIDTH_DXA)
    section.page_height = Twips(15840)
    section.left_margin = Twips(MARGIN_DXA)
    section.right_margin = Twips(MARGIN_DXA)
    section.top_margin = Twips(MARGIN_DXA)
    section.bottom_margin = Twips(MARGIN_DXA)

    # default paragraph font
    style = doc.styles['Normal']
    font = style.font
    font.name = FONT_BODY
    font.size = SZ_BODY
    font.color.rgb = C_BODY

    return doc


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------


def _patch_settings(docx_path: Path):
    """Fix python-docx bug: w:zoom missing required w:percent attribute."""
    import os
    import zipfile

    tmp = docx_path.with_suffix('.tmp.docx')
    with (
        zipfile.ZipFile(docx_path, 'r') as zin,
        zipfile.ZipFile(tmp, 'w', zipfile.ZIP_DEFLATED) as zout,
    ):
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == 'word/settings.xml':
                data = data.replace(b'<w:zoom w:val="bestFit"/>', b'<w:zoom w:percent="100"/>')
            zout.writestr(item, data)
    os.replace(tmp, docx_path)


def convert(md_path: Path, docx_path: Path):
    text = md_path.read_text(encoding='utf-8')

    try:
        md = mistune.create_markdown(renderer='ast', plugins=['table'])
    except TypeError:
        # older/newer Mistune builds may expose plugin registration slightly differently
        md = mistune.create_markdown(renderer='ast')
    ast = md(text)

    doc = _setup_document()
    renderer = DocxRenderer(doc)
    renderer.render_all(ast)

    doc.save(str(docx_path))
    _patch_settings(docx_path)
    print(f'written: {docx_path}')


def main():
    script_dir = Path(__file__).parent

    md_path = Path(sys.argv[1]) if len(sys.argv) > 1 else script_dir / 'cheatsheet.md'
    docx_path = (
        Path(sys.argv[2]) if len(sys.argv) > 2 else script_dir / 'rest_fetcher_cheatsheet.docx'
    )

    if not md_path.exists():
        print(f'error: {md_path} not found', file=sys.stderr)
        sys.exit(1)

    convert(md_path, docx_path)


if __name__ == '__main__':
    main()
